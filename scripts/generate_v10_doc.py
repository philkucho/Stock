"""v10 종목 선정 설명서 → Word 파일 생성.

실행: venv/Scripts/python.exe scripts/generate_v10_doc.py
출력: docs/v10_종목선정_설명서.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ───── 스타일 헬퍼 ─────

def set_korean_font(run, size=10, bold=False, color=None):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rFonts.set(qn("w:ascii"), "맑은 고딕")
    rFonts.set(qn("w:hAnsi"), "맑은 고딕")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sizes = {0: 22, 1: 16, 2: 13, 3: 11}
    set_korean_font(run, size=sizes.get(level, 11), bold=True,
                    color=RGBColor(0x1F, 0x3A, 0x68))
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, size=10, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_korean_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.7 + level * 0.7)
    run = p.runs[-1] if p.runs else p.add_run(text)
    if not p.runs:
        run = p.add_run(text)
    else:
        run.text = text
    set_korean_font(run, size=10)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(doc, headers, rows, col_widths=None, highlight_first=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    # header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_korean_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # 헤더 배경색 (네이비)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F3A68")
        tcPr.append(shd)

    # rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            bold = (c == 0 and highlight_first)
            set_korean_font(run, size=9, bold=bold)

    # col widths
    if col_widths:
        for col_idx, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[col_idx].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table


# ───── 본문 ─────

def build_doc():
    doc = Document()

    # 페이지 마진
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # 표지 제목
    add_heading(doc, "v10 종목 선정 설명서", level=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("미국 주식 자동매매 시스템 — Integrated v10")
    set_korean_font(run, size=12, color=RGBColor(0x55, 0x55, 0x55))
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("작성일: 2026-05-17")
    set_korean_font(run2, size=10, color=RGBColor(0x80, 0x80, 0x80))

    doc.add_paragraph()

    # ─── 0. 한눈에 보기 ───
    add_heading(doc, "한눈에 요약", level=1)
    add_para(doc,
        "이 시스템은 두 가지 매매 모드를 가집니다. 각자 다른 시간대·다른 보유기간을 노립니다.",
        size=11)
    add_table(
        doc,
        headers=["모드", "보유기간", "발송시간", "엔트리 결정", "발송 조건"],
        rows=[
            ("v10 (스윙)",   "수일~수주",  "09:30 ET",  "사용자 직접 입력 (user_fixed)",
             "사용자가 매매 Plan 페이지에 추가한 종목만 그대로 발송"),
            ("orb_auto (단타)", "당일 청산", "09:45 ET",
             "ORB 자동 (스캐너 워치리스트)",
             "스캐너가 09:25 watchlist 산출 → 09:45 ORB+VWAP+RVOL 4-pass 통과만 발송"),
            ("Advisor (AI 자문)", "장중 수시", "매 시간 + 트리거",
             "Gemini LLM 판단 (Telegram 승인)",
             "보유 종목/발송 plan 대상으로 LLM이 enter/add/trim/exit/hold 추천 → 사용자 승인 즉시 실행"),
        ],
        col_widths=[3.0, 2.5, 2.5, 4.5, 6.0],
    )
    add_para(doc,
        "세 모드 모두 v10이 종목 풀의 출발점입니다. v10이 후보를 골라주고, "
        "스윙은 사용자가 그 중 직접 선택해서 추가 / 단타는 시스템이 5개 watchlist를 만들어 09:45에 자동 평가 / "
        "Advisor는 발송된 종목을 장중 내내 모니터링하면서 LLM이 추가/축소/청산을 제안.",
        size=10)
    add_para(doc, "한 줄 철학: \"품질 좋고 + 안전하고 + 지금 사기 좋은 종목\"만 추천한다.",
             size=11, bold=True, color=RGBColor(0xC0, 0x39, 0x2B))

    # ─── Part I 헤더 ───
    add_heading(doc, "Part I — v10 (스윙 매매)", level=0)
    add_para(doc,
        "수일~수주 보유하는 추세 추종(swing) 종목 선정 로직. "
        "매매 Plan 페이지에 표시되며, 사용자가 직접 추가한 종목만 09:30에 발송됩니다.",
        size=10, color=RGBColor(0x55, 0x55, 0x55))

    # ─── 1. 전체 흐름 ───
    add_heading(doc, "1. 전체 선정 흐름 (5단계)", level=1)
    add_para(doc, "매일 아침 장 시작 전 자동으로 다음 순서로 진행됩니다:", size=10)

    steps = [
        ("1단계 — 후보 모집",
         "두 시스템에서 후보를 가져옵니다. v3에서 상위 15개, scanner에서 상위 30개. "
         "중복은 한 번만 세고, 합치면 약 30~40개 종목 풀이 만들어집니다."),
        ("2단계 — 시장 상황 점검",
         "VIX(공포지수), SPY/QQQ 추세, 최근 손실 등으로 시장 모드를 판정합니다. "
         "방어 모드면 아예 빈 리스트를 반환 — 그날은 매수 안 함."),
        ("3단계 — 종목별 안전 검사 (필터)",
         "각 후보 종목에 대해 9가지 위험 검사를 합니다. 하나라도 걸리면 탈락. "
         "예: 7일 내 실적 발표 예정, 최근 30일 reject 2회+, RSI 과열 등."),
        ("4단계 — 점수 계산",
         "필터를 통과한 종목에 대해 v10 합성 점수를 계산합니다. "
         "기본 점수 + 9가지 보너스 점수 + 2가지 배수(곱셈)."),
        ("5단계 — 섹터 분산 + 상위 N개 선정",
         "동일 섹터가 3개 이상이면 페널티를 주고 다시 정렬. 최종 상위 5개가 매매 Plan으로 발송됨."),
    ]
    for title, desc in steps:
        add_para(doc, title, size=11, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
        add_para(doc, "  " + desc, size=10)

    # ─── 2. 후보 풀 ───
    doc.add_page_break()
    add_heading(doc, "2. 후보 풀 — 어디서 종목이 오나?", level=1)
    add_para(doc,
        "v10은 자체적으로 종목을 발굴하지 않습니다. 두 개의 사전 시스템에서 후보를 받아 그 위에 한 번 더 필터링·점수화를 합니다.",
        size=10)

    add_heading(doc, "v3 — 품질 중심 (Quality)", level=2)
    add_para(doc,
        "Minervini 스타일의 '추세 우량주' 발굴. 12개 하드 게이트를 모두 통과한 종목 중 0~100점 만점 점수 상위 15개. "
        "어떤 종목이 통과할까?",
        size=10)
    add_bullet(doc, "장기 추세가 상승 (Stage 2 Trend Template — 200일/150일/50일 이동평균선 정렬)")
    add_bullet(doc, "최근 30일 평균 거래대금 ≥ $20M (충분한 유동성)")
    add_bullet(doc, "주가 $5~$500, 시가총액·float 적정")
    add_bullet(doc, "스프레드 0.3% 이하 (체결 비용 낮음)")
    add_bullet(doc, "변동성(ATR%) 1.5%~12% (너무 죽지도, 너무 과격하지도 않음)")
    add_bullet(doc, "당일 실적 발표 아님")
    add_bullet(doc, "RS Rating(상대강도)·다중 시간프레임 모멘텀·압축/팽창 패턴 등 5개 블록 평가")

    add_heading(doc, "scanner — 거래량+모멘텀 중심 (Volume)", level=2)
    add_para(doc,
        "당일 RVOL(상대거래량) 폭증, 20일 신고가 돌파, 카탈리스트(뉴스/실적) 등으로 '오늘 움직이는' 종목 30개를 추립니다. "
        "v3가 '평소에도 좋은 종목'을 본다면 scanner는 '지금 움직이는 종목'을 봅니다.",
        size=10)

    # ─── 3. 필터 ───
    doc.add_page_break()
    add_heading(doc, "3. 안전 필터 — 무엇을 걸러내나?", level=1)
    add_para(doc,
        "후보 풀에 올라온 종목이라도 다음 조건 중 하나라도 걸리면 탈락합니다. "
        "잘못된 종목을 사는 것보다 안 사는 게 낫다는 철학.",
        size=10)

    filters = [
        ("시장 방어 모드", "regime_score < 7 또는 VIX 급등 + 추세 약함 → 그날은 매수 자체를 안 함."),
        ("자동 블랙리스트", "최근 30일 내 outcome alpha < -1%인 reject가 2번 이상 누적된 종목 → 시스템이 스스로 거부."),
        ("실적 발표 임박", "0~7일 내 실적 발표 예정 → 변동성 위험 회피. (단, 발표 직후 종목은 보너스 ×1.25)"),
        ("최근 피드백 부정", "최근 30일 시간가중 평균 alpha < -1% → 자동 reject. 1일내 4배, 5일내 2배, 15일내 1배 가중치."),
        ("RSI 과열", "RSI grade='bad' (climax 또는 bearish divergence) → 탈락."),
        ("Gap-and-fail 위험", "시초가가 전일 고점 위로 갭 떴는데 곧바로 무너지는 패턴 → 탈락."),
        ("VIX > 25 + 골든 셋업 부재", "변동성 큰 장에서는 압축+팽창이 둘 다 있는 종목만 통과."),
        ("Drawdown 모드", "최근 10일 integrated picks 중 -5% 이하 outcome 2개+ → 보수 모드. 압축 또는 RSI good 또는 anchored VWAP 위만 통과."),
        ("Scanner 단독 진입 시 신호 3개 이상 필수", "Tier 2(scanner only)는 9개 신호 중 3개 미만이면 탈락."),
    ]
    add_table(
        doc,
        headers=["필터", "탈락 조건"],
        rows=filters,
        col_widths=[5.0, 12.0],
    )

    # ─── 4. 점수 계산 ───
    doc.add_page_break()
    add_heading(doc, "4. 점수 계산 — v10 합성 점수는 어떻게?", level=1)
    add_para(doc,
        "필터를 통과한 종목에 대해 다음과 같이 점수를 계산합니다. "
        "기본 점수에 보너스를 더한 뒤 두 가지 배수(곱셈)를 적용합니다.",
        size=10)

    add_para(doc, "공식: 점수 = (기본점수 + 보너스 합계) × earnings_multiplier × super_multiplier",
             size=10, bold=True)

    add_heading(doc, "4-1. 기본 점수 — Tier 별로 다름", level=2)
    add_para(doc, "v3에서 온 종목 (Tier 1):", size=10, bold=True)
    add_table(
        doc,
        headers=["항목", "만점", "설명"],
        rows=[
            ("v3_norm", "50점", "v3 자체 점수의 1.3제곱 (강한 종목 강조)"),
            ("ce_norm", "20점+α", "압축+팽창 패턴 점수 × regime 보너스 + golden setup 10점 + stage2 5점"),
            ("ol_norm", "8점", "시가 위치 (open location)"),
            ("rsi_norm", "5점", "RSI 구조 (good=5, ok=3, neutral=0)"),
        ],
        col_widths=[3.5, 2.5, 11.0],
    )
    add_para(doc, "scanner 단독 종목 (Tier 2): scanner_norm 25점, ce 25점, ol 12점, rsi 8점 — 압축+팽창을 더 중시.",
             size=10)

    add_heading(doc, "4-2. 보너스 점수 (모두 더함)", level=2)
    bonuses = [
        ("sector_bonus", "+10", "반도체·테크 등 우선 섹터"),
        ("sector_mom_bonus", "+5", "섹터 ETF 5일 모멘텀이 SPY보다 +0.5% 강함"),
        ("confluence_bonus", "+15", "scanner도 4점 이상 (Tier 1만)"),
        ("stage2_bonus", "+8 / +12", "Minervini Stage 2 통과 (Tier 1 / Tier 2)"),
        ("streak_bonus", "+10~15", "v3에서 연속 3회 이상 추천된 종목"),
        ("obv_bonus", "+5", "OBV(누적 매수세) 우상향"),
        ("mom_accel_bonus", "+5", "1개월 상대강도 > 3개월 상대강도 (모멘텀 가속)"),
        ("avwap_bonus", "+5", "주가가 Anchored VWAP 위에 있음"),
        ("feedback_bonus", "±12", "최근 30일 시간가중 alpha × 5 (성공시 +, 실패시 −)"),
    ]
    add_table(
        doc,
        headers=["보너스", "점수", "조건"],
        rows=bonuses,
        col_widths=[4.5, 2.0, 10.5],
    )

    add_heading(doc, "4-3. 배수 (곱셈)", level=2)
    add_para(doc, "earnings_multiplier = 1.25 (PEAD 알파)",
             size=10, bold=True)
    add_para(doc,
        "실적 발표 직후(post-earnings) 종목은 합성 점수에 ×1.25 적용. "
        "발표 전은 회피(필터)하지만 발표 후는 통계적으로 +2.5% 알파가 검증됨.",
        size=10)

    add_para(doc, "super_multiplier = 1.3 (5중 합의)",
             size=10, bold=True)
    add_para(doc,
        "다음 5개 강한 신호가 동시에 충족되면 ×1.3 부스트:",
        size=10)
    add_bullet(doc, "v3 통과 (Tier 1) 또는 scanner ≥ 4점 (Tier 2)")
    add_bullet(doc, "압축(compression) 동시 발생")
    add_bullet(doc, "팽창(expansion) 동시 발생 — 둘이 같이 있어야 'golden setup'")
    add_bullet(doc, "v3 연속 추천 3회 이상")
    add_bullet(doc, "RSI grade='good'")

    add_heading(doc, "4-4. 섹터 분산 페널티", level=2)
    add_para(doc,
        "정렬 후, 동일 섹터의 3번째 종목부터 (순서-1) × 5점 감점. 한 섹터에 몰리는 것 방지.",
        size=10)

    # ─── 5. 실제 샘플 ───
    doc.add_page_break()
    add_heading(doc, "5. 실제 샘플 — 2026-05-15 선정 종목", level=1)
    add_para(doc,
        "마지막 거래일(2026-05-15, 금요일)에 v10이 실제로 뽑은 4개 종목을 봅시다. "
        "(이 날은 일부 종목이 필터 탈락해서 5개가 아닌 4개로 줄어듦)",
        size=10)

    add_table(
        doc,
        headers=["순위", "종목", "섹터", "점수", "경로"],
        rows=[
            ("1위", "AMAT", "Technology (반도체장비)", "88.09", "v3_priority"),
            ("2위", "TXN",  "Technology (반도체)",     "58.32", "v3_priority"),
            ("3위", "CSCO", "Technology (네트워크)",   "48.44", "v3_priority"),
            ("4위", "LRCX", "Technology (반도체장비)", "77.44", "v9_fallback"),
        ],
        col_widths=[2.0, 2.5, 5.5, 2.0, 5.0],
    )

    add_heading(doc, "AMAT (Applied Materials) — 1위, 88.09점", level=2)
    add_para(doc, "왜 1등이 되었나?", size=10, bold=True)
    amat_rows = [
        ("v3 기본 점수", "~36점", "v3 score 52.59 → (0.5259)^1.3 × 50"),
        ("ce_norm", "0점", "compression/expansion 없음 (이번엔 보너스 없음)"),
        ("open_location", "~6.4점", "open_location_score 4.0 → 4/5 × 8"),
        ("rsi_norm", "5점", "RSI grade=good → 만점"),
        ("sector_bonus", "+10", "Technology (우선 섹터)"),
        ("sector_mom_bonus", "+5", "반도체 ETF 모멘텀 +4.12% (SPY 대비)"),
        ("stage2_bonus", "+8", "Stage 2 trend template 통과"),
        ("streak_bonus", "+10", "v3에서 4일 연속 추천 (streak_count=4)"),
        ("obv_bonus", "+5", "OBV 우상향 (누적 매수세)"),
        ("avwap_bonus", "+5", "anchored VWAP 위"),
        ("feedback_bonus", "+12", "최근 30일 평균 alpha +3.32% (최대치)"),
        ("earnings_multiplier", "×1.0", "post-earnings 아님"),
        ("super_multiplier", "×1.0", "압축+팽창 없어서 super 미발동"),
        ("최종", "88.09", "기본+보너스 합계"),
    ]
    add_table(doc, ["항목", "기여", "설명"], amat_rows, col_widths=[5.0, 2.5, 9.5])

    add_para(doc, "해석: 압축/팽창 패턴은 없지만 다른 모든 신호가 강함. "
                  "'추세 + 섹터 + 연속 추천 + 최근 성과 피드백'이 동시에 잘 맞아 88점. "
                  "5중 합의(super)는 아니지만 안정적 setup.", size=10)

    add_heading(doc, "TXN (Texas Instruments) — 2위, 58.32점", level=2)
    add_para(doc, "Technology 섹터지만 우선 섹터 보너스가 안 붙음 — 왜?", size=10)
    add_para(doc, "  score_meta를 보면 sector_bonus=0. AMAT/LRCX와 달리 'Semiconductor'가 아닌 일반 Technology로 분류돼서 "
                  "PRIORITY_SECTORS 매칭이 안 됨. 또 streak_count=0이라 streak 보너스도 없음. "
                  "그래도 momentum_accel=true(+5), feedback +9.56, obv·avwap·stage2 보너스로 58점.",
             size=10)

    add_heading(doc, "CSCO (Cisco) — 3위, 48.44점", level=2)
    add_para(doc, "TXN과 비슷한 구조지만 open_location_score=0 (시초가가 피벗/전일고에서 멀리 벗어남). "
                  "또 동일 Technology 섹터 3번째라 diversification_penalty -5점 발생. 결국 48점.",
             size=10)

    add_heading(doc, "LRCX (Lam Research) — 4위, 77.44점 (v9 fallback)", level=2)
    add_para(doc, "이 종목은 source='v9_fallback'. v10 알고리즘 산출 실패 시 v9 결과로 대체. "
                  "streak_count=5 (5일 연속), feedback_avg_alpha 4.05%, OBV·AVWAP·Stage 2 모두 통과. "
                  "AMAT와 비슷한 setup이지만 v9이라 점수 체계가 약간 다름.",
             size=10)

    # ─── 6. 실제 성과 ───
    doc.add_page_break()
    add_heading(doc, "6. 그래서 정말 잘 맞나? — 실제 outcome", level=1)
    add_para(doc,
        "최근 v10이 추천한 종목들의 실제 5일/10일 수익률입니다. "
        "alpha는 SPY 대비 초과수익(%). 빨간 글씨는 손실, 굵은 글씨는 알파 +10% 초과.",
        size=10)

    outcomes = [
        ("2026-04-23", "MU",   "5d", "+9.33%", "+7.94%"),
        ("2026-04-23", "MU",   "10d", "+50.59%", "+46.81%"),
        ("2026-04-24", "MU",   "5d", "+12.91%", "+12.23%"),
        ("2026-04-24", "MU",   "10d", "+55.78%", "+52.11%"),
        ("2026-04-24", "ON",   "10d", "+9.28%", "+5.62%"),
        ("2026-04-27", "MRVL", "5d", "+14.09%", "+12.41%"),
        ("2026-04-27", "ON",   "10d", "+9.58%", "+5.87%"),
        ("2026-04-28", "VRT",  "5d", "+16.89%", "+13.68%"),
        ("2026-04-28", "VRT",  "10d", "+20.50%", "+16.09%"),
        ("2026-04-29", "VRT",  "10d", "+19.91%", "+15.22%"),
        ("2026-04-29", "ON",   "10d", "+18.11%", "+13.42%"),
        ("2026-04-30", "VRT",  "10d", "+12.19%", "+9.70%"),
        ("2026-04-30", "ON",   "10d", "+12.11%", "+9.63%"),
        ("2026-04-30", "TXN",  "10d", "+8.08%",  "+5.59%"),
        ("2026-05-05", "VRT",  "5d", "+5.24%",  "+3.30%"),
        ("2026-05-06", "VRT",  "5d", "+4.92%",  "+3.14%"),
        ("2026-05-07", "VRT",  "5d", "+6.05%",  "+5.47%"),
    ]
    add_table(
        doc,
        headers=["pick_date", "종목", "기간", "수익률", "alpha (vs SPY)"],
        rows=outcomes,
        col_widths=[3.0, 2.5, 2.0, 3.5, 4.0],
    )
    add_para(doc,
        "관찰: 4월말~5월초 picks(VRT/MU/MRVL/ON)는 10일 수익률이 +9~+55% 범위. "
        "MU는 실적 발표 직후 PEAD 알파가 폭발(10d +55%). VRT는 5일 연속 추천 종목 중 streak 효과 큼.",
        size=10)

    add_heading(doc, "[[project_integrated_v10]] 백테스트 결과", level=2)
    add_bullet(doc, "60일 backfill — 10일 평균 alpha: +14.88%")
    add_bullet(doc, "Win rate (10일 기준): 93%")
    add_bullet(doc, "Sharpe ratio: 6.58")
    add_bullet(doc, "v9 대비 개선: 5일 plateau 돌파 (7.04% → 14.88%)")

    # ─── 7. 자주 묻는 질문 ───
    doc.add_page_break()
    add_heading(doc, "7. 자주 묻는 질문", level=1)

    qa = [
        ("Q. 왜 매일 5개가 아니라 4개만 나왔지?",
         "A. 12개 게이트 + 9개 필터를 통과한 종목이 5개 안 되는 날이 있음. "
         "강제로 5개 채우지 않음 — 부족하면 그대로 4개, 3개로 발송. 잘못된 종목 사느니 비워두는 게 낫다는 원칙."),
        ("Q. v3와 scanner 둘 다 안 뽑은 종목은 절대 안 들어가나?",
         "A. 맞음. v10은 후보 발굴을 안 함. 따라서 v3 top 15에도 없고 scanner top 30에도 없는 종목은 v10에서 절대 추천되지 않음. "
         "이게 단점이기도 한데, 두 시스템의 cross-validation 효과로 잘못된 추천을 줄이는 게 더 중요하다고 판단."),
        ("Q. 점수가 88점이면 무조건 좋은가?",
         "A. 절대 점수는 의미 적음. 같은 날 다른 종목 대비 상대적으로 높은지가 중요. "
         "예: 시장이 약하면 모든 종목 점수가 낮아지고, 강하면 다 높아짐. 매일 top 5로 줄세우기 위한 ranking signal."),
        ("Q. AI(Advisor)가 추가로 거를 수 있나?",
         "A. 가능. v10 picks가 1차 후보가 되고, Gemini가 regime/포지션/뉴스/최근 성과를 보고 그 중 일부만 추천하거나 모두 reject할 수 있음. "
         "[[project_advisor_v1]] 참고."),
        ("Q. 며칠 동안 똑같은 종목이 계속 나오면 어떡하지?",
         "A. 그건 사실 좋은 신호 — streak_bonus 적용 대상. v3에서 3일 연속 추천되면 +10, 5일+ 연속이면 +15. "
         "강한 추세 종목은 계속 top에 올라옴. 반대로 최근 outcome이 나쁘면 feedback_bonus가 음수가 돼서 점차 빠짐."),
        ("Q. 종목 매수 가격(entry)과 손절가(stop)는 어떻게 정하나?",
         "A. v10 자체는 ranking만. entry/stop/1R/2R은 trading._build_picks의 compute_pick_metadata에서 별도 산출 — "
         "60일 일봉 + 5일 5분봉으로 pivot, ATR 기반 stop, 1R/2R 자동 계산. 자세한 건 stage2_daily_picks.py 참조."),
    ]
    for q, a in qa:
        add_para(doc, q, size=10, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
        add_para(doc, a, size=10)

    # ─── Part II 헤더 ───
    doc.add_page_break()
    add_heading(doc, "Part II — orb_auto (단타 매매)", level=0)
    add_para(doc,
        "당일 청산을 목표로 하는 단타(intraday) 자동 매매 로직. "
        "스캐너가 09:25에 watchlist를 산출하고, 09:45에 ORB+VWAP+RVOL 4-pass로 자동 평가해서 통과한 종목만 bracket order로 발송합니다. "
        "현재 [[project_hybrid_dispatch]] 정책에 따라 AUTO_CONFIRM_DISPATCH=false (기본 OFF) — 자동 발송은 비활성, 수동 검증/관찰 단계.",
        size=10, color=RGBColor(0x55, 0x55, 0x55))

    # ─── 8. orb_auto 한눈에 ───
    add_heading(doc, "8. orb_auto — 한눈에 요약", level=1)
    add_para(doc,
        "orb_auto는 \"오늘 아침 강하게 움직이는 종목을 09:45 ET 시점에 자동으로 골라 사고, 당일 안에 청산\"을 목표로 합니다. "
        "v10이 매수해도 좋은 후보를 미리 가져다놓고(09:25), 시장이 실제로 그 종목을 사는지(브레이크아웃) 15분간 관찰한 뒤(09:30~09:44) "
        "정말 강하게 움직이는 종목만 09:45에 매수합니다.",
        size=11)
    add_para(doc, "한 줄 철학: \"가설(후보)을 세우고, 시장이 인정한 종목만 산다\"",
             size=11, bold=True, color=RGBColor(0xC0, 0x39, 0x2B))

    # ─── 9. 전체 흐름 ───
    add_heading(doc, "9. orb_auto 전체 흐름 (2단계 시간 분리)", level=1)
    add_para(doc,
        "스윙(v10)은 한 번에 끝나지만, orb_auto는 두 단계로 나뉩니다. 시간이 다릅니다.",
        size=10)

    add_heading(doc, "Phase 4 — Preopen (09:25 ET): 워치리스트 생성", level=2)
    phase4_steps = [
        ("1. v10 candidate pool 가져오기",
         "v10에서 top 10을 받음 (\"품질+안전 합격 후보\"). 단, top 10이라서 swing은 5개만 사용하는 것과 달리 더 넓게."),
        ("2. 프리마켓 데이터 수집",
         "각 종목의 프리마켓 시가/종가/갭/RVOL/스프레드를 yfinance로 실시간 fetch."),
        ("3. 하드 스킵",
         "갭 > +10% (과열) / 갭 < -3% (약세 갭) / 스프레드 > 1.5% (유동성 부족) → 탈락."),
        ("4. 카탈리스트 확인",
         "어닝/뉴스/업그레이드 등 catalyst aggregator로 점수 추가."),
        ("5. 5-Model 합성 점수",
         "v10 점수 + 갭 + RVOL + catalyst 4개를 0~1 정규화 후 가중합. 섹터 분산 페널티 적용."),
        ("6. Top 5 → trade_plans 테이블에 watchlist 상태로 저장",
         "dispatch_mode='orb_auto', confirm_status='watchlist'. provisional entry/stop/1R/2R도 미리 계산해서 넣어둠 (09:45에 덮어쓰기)."),
    ]
    for title, desc in phase4_steps:
        add_para(doc, title, size=11, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
        add_para(doc, "  " + desc, size=10)

    add_heading(doc, "Phase 5 — Confirm (09:45 ET): ORB 평가 + 발송", level=2)
    phase5_steps = [
        ("1. watchlist 로드",
         "dispatch_mode='orb_auto' AND confirm_status='watchlist' 인 plan만 가져옴. user_fixed는 09:30 cron이 별도 처리."),
        ("2. 09:30~09:44 1분봉 fetch (15개 bar)",
         "yfinance 1m bars 다운로드. 마지막 bar가 30분 이상 오래되면 stale로 취급, 스킵."),
        ("3. ORB+VWAP+RVOL 4-pass 평가",
         "다음 4개 조건을 모두 충족해야 통과. 하나라도 빠지면 confirm_status='failed'로 마킹."),
        ("4. ORB 기반 entry/stop/target 재계산",
         "entry = OR high + $0.05 / stop = max(OR low, session VWAP) / target = entry+R, entry+2R. R/entry < 0.3%면 거부."),
        ("5. Position sizing",
         "equity × 0.3% (INTRADAY_RISK_PCT) / R × regime_mult × gap_penalty. 갭>5%면 ×0.7, 갭>10%면 ×0."),
        ("6. 안전장치 통과 후 bracket order 발송",
         "regime defensive, account blocked, position cap 5, sector cap 2, daily loss limit 등 9중 안전장치. 최대 top 3 종목만 발송."),
        ("7. plan 상태 업데이트",
         "confirm_status='sent' + broker_order_ids 저장. ORB 데이터(orb_high/orb_low/session_vwap/intraday_rvol)도 plan에 저장."),
    ]
    for title, desc in phase5_steps:
        add_para(doc, title, size=11, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
        add_para(doc, "  " + desc, size=10)

    # ─── 10. 후보 풀 ───
    doc.add_page_break()
    add_heading(doc, "10. orb_auto 후보 풀 — 5-Model Stack", level=1)
    add_para(doc,
        "watchlist 산출 단계(09:25)에서는 다음 5개 신호를 stack합니다. v10은 base이고, 그 위에 프리마켓 신호 4개를 곱하는 구조.",
        size=10)
    stack = [
        ("① v10 base",        "60점", "v10 종목 점수 / 250 → 0~1 정규화 × 60. 어제 종가 셋업 강도 반영."),
        ("② Premarket Gap",   "20점", "오늘 시초가 갭 / 5% → 0~1 saturate × 20. +5% 갭이면 만점."),
        ("③ Premarket RVOL",  "15점", "(RVOL - 1) / 2 → 0~1 saturate × 15. RVOL 3배면 만점."),
        ("④ Catalyst",        "15점", "catalyst.score / 30 (KIND_SCORE max) × 15. 어닝/뉴스 보너스."),
        ("(⑤ Regime)",        "—",    "Regime score 기반 sizing multiplier. 점수 자체에는 안 들어가고 Phase 5 sizing에서 사용."),
    ]
    add_table(doc, ["모델", "만점", "설명"], stack, col_widths=[4.0, 1.5, 11.5])

    add_para(doc, "총점 = ①+②+③+④ (최대 110점). 정렬 후 섹터 분산 페널티(2개 초과 시 (n-1)×5) → Top 5.",
             size=10, bold=True)

    add_heading(doc, "Hard skip — 다음 중 하나라도 걸리면 탈락", level=2)
    hard_skips = [
        ("갭 +10% 초과",   "INTRA_GAP_HARD_SKIP. 과열 상태에서 진입은 fade 위험 큼."),
        ("갭 -3% 미만",    "INTRA_GAP_LOW_SKIP. 약세 갭에서 long은 위험."),
        ("스프레드 1.5% 초과", "INTRA_SPREAD_MAX. 체결 비용으로 R 잠식."),
        ("프리마켓/전일종가 누락", "데이터 fetch 실패 시 안전하게 skip."),
        ("provisional R ≤ 0",  "entry < stop이면 invalid."),
    ]
    add_table(doc, ["스킵 사유", "이유"], hard_skips, col_widths=[5.0, 12.0])

    # ─── 11. ORB 4-pass ───
    doc.add_page_break()
    add_heading(doc, "11. ORB 4-pass — 09:45 ET 실제 발송 조건", level=1)
    add_para(doc,
        "watchlist에 올라왔다고 발송되는 게 아닙니다. 09:30~09:44 첫 15분 동안 시장이 \"이 종목 진짜 사고 있다\"는 신호를 4개 모두 보여줘야 발송. "
        "하나라도 빠지면 그 종목은 confirm_status='failed'로 끝.",
        size=10)

    orb_pass = [
        ("① ORB 돌파 (pass_orb)",
         "현재가 > Opening Range High × (1 + 0.1%)",
         "처음 15분의 최고가를 0.1% 이상 뚫어야 함. 그냥 high를 찍는 게 아니라 명확하게 돌파해야."),
        ("② VWAP 위 (pass_vwap)",
         "현재가 > session VWAP",
         "기관 평균매수가(VWAP) 위에서 거래되어야 함. VWAP 아래는 매도 우위라 단타 long 불리."),
        ("③ Intraday RVOL (pass_rvol)",
         "오늘 첫 15분 거래량 ≥ 직전 20일 동시간대 거래량 중위값 × 1.5",
         "거래량 폭증이 있어야 함. 평소 같은 시간대 대비 1.5배 이상."),
        ("④ OR Range 충분 (pass_range)",
         "(OR high - OR low) / 현재가 ≥ 0.5%",
         "첫 15분의 high-low 폭이 너무 좁으면 fake breakout 위험 → 0.5% 이상 변동 필수."),
    ]
    add_table(doc, ["조건", "수식", "왜?"], orb_pass, col_widths=[4.5, 5.0, 7.5])

    add_heading(doc, "통과 시 entry/stop/target 결정", level=2)
    add_para(doc, "ORB 데이터로 매매 가격 4개를 재산정합니다 (Phase 4의 provisional 값을 덮어씀):",
             size=10)
    add_bullet(doc, "entry = OR high + $0.05  (penny offset, breakout 확인용)")
    add_bullet(doc, "stop  = max(OR low, session VWAP)  (둘 중 높은 쪽 = 더 보수적)")
    add_bullet(doc, "R     = entry - stop")
    add_bullet(doc, "target_1r = entry + R")
    add_bullet(doc, "target_2r = entry + 2R")
    add_bullet(doc, "R/entry < 0.3% 이면 stop이 너무 가까워서 거부 (slippage 위험)")

    # ─── 12. 점수 계산 ───
    add_heading(doc, "12. Position Size 계산", level=1)
    add_para(doc, "통과한 종목 각각의 매수 수량은 다음과 같이 산정:", size=10)
    add_para(doc, "qty = floor(equity × 0.3% / R) × regime_mult × gap_penalty",
             size=11, bold=True)
    add_table(
        doc,
        headers=["변수", "값", "의미"],
        rows=[
            ("equity", "계좌 자본금", "Alpaca paper account 현재 자본"),
            ("0.3%", "INTRADAY_RISK_PCT", "1트레이드당 위험 비율. 100만원 계좌면 3000원 손실 한도."),
            ("R", "entry - stop", "주당 위험금액"),
            ("regime_mult", "0.5~1.2", "방어 모드면 0.5, 공격 모드면 1.2"),
            ("gap_penalty", "0~1.0", "갭 +5%~+10% → 0.7, 갭 +10%+ → 0 (skip)"),
        ],
        col_widths=[3.0, 3.0, 11.0],
    )
    add_para(doc, "발송 직전 9개 안전장치: regime defensive / account blocked / daily loss limit / position cap 5 / "
                  "sector cap 2 / 이미 보유 / pending 2개 이상 / buying power 부족 / R invalid. "
                  "통과한 종목 중 composite 점수 상위 3개만 실제 발송 (INTRADAY_PICK_CAP=3).",
             size=10)

    # ─── 13. 실제 샘플 + 백테스트 ───
    doc.add_page_break()
    add_heading(doc, "13. 실제 샘플 — ORB 백테스트 결과 (2026-03-17 ~ 2026-05-14)", level=1)
    add_para(doc,
        "현재 운영 DB에 실제 orb_auto 발송 이력은 없습니다 (AUTO_CONFIRM_DISPATCH=false). "
        "대신 백테스트로 같은 알고리즘을 과거 데이터에 돌린 결과는 있습니다:",
        size=10)

    add_table(
        doc,
        headers=["지표", "10종목 60일", "AVGO+TXN 2종목 단독"],
        rows=[
            ("거래 횟수",       "6건",        "2건"),
            ("승률",            "0%",         "0%"),
            ("평균 R-multiple", "-1.0R",      "-1.0R"),
            ("평균 수익률",     "-0.69%",     "-0.40%"),
            ("총 손익",         "-$1,233.70", "-$265.92"),
            ("Sharpe (연환산)", "-15.25",     "-70.80"),
            ("청산 사유",       "전부 stop",  "전부 stop"),
        ],
        col_widths=[4.5, 4.5, 6.5],
    )

    add_heading(doc, "샘플 거래 — 6건 전체 stop-out 내역", level=2)
    add_table(
        doc,
        headers=["날짜", "종목", "entry", "stop", "1R 목표", "PnL"],
        rows=[
            ("2026-03-20", "AAPL", "$248.78", "$247.98", "$249.58", "-$79.92"),
            ("2026-04-07", "AAPL", "$256.25", "$250.73", "$261.77", "-$552.18"),
            ("2026-04-14", "AVGO", "$381.84", "$380.55", "$383.13", "-$128.54"),
            ("2026-04-16", "MSFT", "$420.05", "$418.13", "$421.97", "-$192.23"),
            ("2026-04-24", "AMD",  "$347.52", "$346.09", "$348.95", "-$143.45"),
            ("2026-05-11", "TXN",  "$296.35", "$294.97", "$297.72", "-$137.38"),
        ],
        col_widths=[3.0, 2.0, 2.5, 2.5, 2.5, 2.5],
    )

    add_heading(doc, "왜 다 stop-out 됐나? — 진단", level=2)
    diag_para = add_para(doc,
        "백테스트 진단 데이터(backtests/results/intraday_orb_diag.json)를 보면 패턴이 보입니다.",
        size=10)
    add_bullet(doc, "AAPL 2026-03-20: OR high $248.73 돌파했지만 09:45 시점 가격 $247.33 — 진입 직후 곧장 무너짐")
    add_bullet(doc, "RVOL은 4.0배(매우 강함)이고 OR Range도 0.85%로 충분 — 4-pass 통과는 정상")
    add_bullet(doc, "공통: stop이 entry 가까이($1~2 차이) 있어서 변동성에 쉽게 hit. 1R 작은데 시장 진폭이 더 큼")
    add_bullet(doc, "결론: ORB는 정상 작동하지만 \"돌파 후 fade\" 패턴에서 stop loss로 즉시 종료. "
                    "기존 swing(v10)이 +14~55% 알파를 내는 동안 단타는 손실 누적 — 그래서 현재 자동 발송 OFF.")

    add_para(doc, "운영 정책 (2026-05-13):",
             size=10, bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
    add_para(doc,
        "사용자 결정 — \"plan에 없는 것은 주문하지 말아.\" "
        "→ AUTO_CONFIRM_DISPATCH=false. orb_auto는 watchlist 산출 + ORB 평가까지만 동작하고 "
        "bracket order 발송은 비활성. 매매 Plan UI에서 사용자가 보고 판단만 한다.",
        size=10)
    add_para(doc, "다시 켜려면 .env에 AUTO_CONFIRM_DISPATCH=true 추가.",
             size=9, color=RGBColor(0x80, 0x80, 0x80))

    # ─── 14. FAQ ───
    doc.add_page_break()
    add_heading(doc, "14. orb_auto 자주 묻는 질문", level=1)
    qa_orb = [
        ("Q. v10과 orb_auto의 종목이 같은가?",
         "A. 출발점은 같음 — 둘 다 v10에서 후보를 받음. 차이는 ① v10 swing은 top 5만, orb_auto는 top 10에서 5개 추림. "
         "② orb_auto는 그 위에 프리마켓 신호(갭/RVOL/catalyst)를 추가로 적용해서 재정렬. "
         "③ 09:45 ORB 평가에서 추가 탈락. 결과적으로 같은 날 v10 swing top 5와 orb_auto sent 종목이 다를 수 있음."),
        ("Q. 09:25에 watchlist가 만들어진 종목은 09:45에 무조건 매수되나?",
         "A. 절대 아님. watchlist는 \"후보 가설\"일 뿐이고, 실제 발송은 4-pass(ORB+VWAP+RVOL+Range) 통과한 종목 중 "
         "상위 3개만. 보통 watchlist 5개 중 1~2개만 실제 발송됨. 4-pass 실패하면 confirm_status='failed'."),
        ("Q. AUTO_CONFIRM_DISPATCH=false인데 왜 watchlist는 계속 만들어지나?",
         "A. 발송만 OFF, 평가는 ON. /trading 화면에서 사용자가 watchlist를 보면서 \"오늘 어떤 종목이 강세인지\" 판단 가능. "
         "필요하면 사용자가 user_fixed 모드로 직접 입력해서 발송. 즉 orb_auto는 \"AI 후보 추천기\"로 동작 중."),
        ("Q. ORB high만 보고 매수하는 것 같은데 너무 단순한 거 아닌가?",
         "A. 단일 신호 아님. 4개를 동시에 봄: ① OR high 돌파(가격) ② VWAP 위(기관 평균) ③ RVOL 1.5배(거래량) "
         "④ OR Range 0.5%+(충분한 변동). 모두 만족해야 \"실제 강한 brearkout\". 거기에 entry/stop도 ORB 기반으로 재계산."),
        ("Q. 백테스트가 6/6 stop이면 알고리즘이 잘못된 것 아닌가?",
         "A. 가능성 있음. 두 가지 가설 — ① ORB 알고리즘 자체가 미국 대형주에서는 안 먹힘 (학계 연구로도 small cap에서 더 잘 통함). "
         "② 백테스트 샘플(6건)이 너무 적음. 더 다양한 종목/기간 필요. "
         "현재는 운영하면서 데이터 더 쌓고 (paper trading) 알고리즘 fine-tune 단계."),
        ("Q. orb_auto의 entry/stop은 v10과 다른가?",
         "A. 매우 다름. v10 swing은 60일 일봉으로 pivot/ATR 기반 entry/stop 산출(보통 R = $2~$10). "
         "orb_auto는 첫 15분 OR high/low 기반(보통 R = $0.50~$3). 단타는 R이 훨씬 작고 빠르게 결판. "
         "그래서 1트레이드당 risk pct도 더 작음 (0.3% vs swing의 0.5~1%)."),
    ]
    for q, a in qa_orb:
        add_para(doc, q, size=10, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
        add_para(doc, a, size=10)

    # ─── Part III 헤더 ───
    doc.add_page_break()
    add_heading(doc, "Part III — Advisor (AI 자문, 매 시간 추천)", level=0)
    add_para(doc,
        "Gemini 2.5 Flash가 장중 내내 보유 종목과 발송된 plan을 모니터링하면서 "
        "추가매수(add) / 부분매도(trim) / 청산(exit) / 신규진입(enter) / 유지(hold)를 추천하는 시스템. "
        "추천은 Telegram inline 키보드로 사용자에게 전달되고, 사용자가 Approve를 누르면 즉시 실행됨. "
        "[[project_advisor_v1]] · 2026-05-15 Approve 즉시 발송 완성.",
        size=10, color=RGBColor(0x55, 0x55, 0x55))

    # ─── 15. Advisor 한눈에 ───
    add_heading(doc, "15. Advisor — 한눈에 요약", level=1)
    add_para(doc,
        "스윙(v10)·단타(orb_auto)가 \"종목을 찾아서 사는\" 시스템이라면, Advisor는 \"이미 산 종목을 어떻게 관리할지\" 결정하는 시스템입니다. "
        "장중 가격/뉴스/거래량 변화에 반응하고, 매 시간 정기적으로 한 번씩 LLM에게 \"지금 이 종목 어떻게 할까?\"를 묻습니다.",
        size=11)
    add_para(doc, "한 줄 철학: \"진입은 알고리즘이, 관리는 AI+사람이.\"",
             size=11, bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
    add_para(doc,
        "현재 ADVISOR_ENABLED=true로 활성화되어 매시간 + 트리거 기반으로 동작 중. "
        "비용 발생 (Gemini API 호출). 사용자 승인 없이는 어떤 주문도 자동 실행되지 않음.",
        size=10)

    # ─── 16. 전체 흐름 ───
    add_heading(doc, "16. Advisor 전체 흐름 (5단계)", level=1)
    advisor_steps = [
        ("1. 모니터 set 구성",
         "오늘 발송된 trade_plan (broker_order_ids 있는 것) + 현재 broker 보유 포지션을 합집합으로 set 구성. "
         "watchlist에만 있고 발송 안 된 종목은 제외 (진입 안 했으므로 add/trim/exit 의미 없음)."),
        ("2. 트리거 평가",
         "두 가지 모드. ① 15분 cron — 가격/뉴스/거래량 3가지 트리거 평가. 하나라도 충족하면 LLM 호출. "
         "② 매 시간 정기 검토 cron — 트리거 무시하고 monitor set 전종목 LLM 호출."),
        ("3. dedupe 확인",
         "같은 종목에 대해 30분 안에 이미 추천이 있으면 skip (LLM 비용 절약)."),
        ("4. LLM 호출 + 결정",
         "각 종목의 컨텍스트(trade_plan, 포지션, 30개 1분봉, 6시간 뉴스, regime, trigger 사유)를 dict로 만들어 Gemini에 전달. "
         "Gemini는 enter/add/trim/exit/hold 중 하나 + confidence + reasoning 반환."),
        ("5. 저장 + 알림 + 승인 실행",
         "confidence ≥ 0.6이면 advisor_recommendations 테이블에 저장 + Telegram inline 키보드(Approve/Reject) 발송. "
         "사용자 Approve → 즉시 broker API로 실행. Reject → 사유 저장. 5분 후 자동 expired."),
    ]
    for title, desc in advisor_steps:
        add_para(doc, title, size=11, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
        add_para(doc, "  " + desc, size=10)

    # ─── 17. 두 가지 모드 ───
    doc.add_page_break()
    add_heading(doc, "17. 두 가지 호출 모드 — 트리거 vs 시간별", level=1)

    add_heading(doc, "모드 A — 트리거 기반 (15분 cron, 09:50~15:55 ET)", level=2)
    add_para(doc,
        "monitor set 종목 각각에 대해 \"지금 비정상적 신호가 있나?\" 평가. "
        "3가지 트리거 중 하나라도 충족하면 LLM 호출. 충족 신호가 없으면 LLM 안 부름 (비용 0).",
        size=10)
    triggers = [
        ("price_spike / price_drop",
         "1분봉 ±2σ 이탈",
         "직전 30개 1분봉의 평균/표준편차 대비 현재 종가가 ±2σ 벗어났을 때. "
         "보유 종목이 갑자기 급등(spike) 또는 stop 근처로 급락(drop)했음을 의미."),
        ("news",
         "신규 뉴스 15분 이내",
         "Finnhub 종목 뉴스 API에서 마지막 15분 안에 발행된 헤드라인이 있으면. "
         "어닝 발표, 규제 뉴스, 업그레이드/다운그레이드 등에 즉각 반응."),
        ("rvol",
         "직전 5분 거래량 ≥ 평소 2배",
         "최근 1분봉의 거래량이 직전 평균의 2배 이상이면. "
         "기관 매수/매도 흔적 — 가격 움직임 전조."),
    ]
    add_table(doc, ["트리거", "조건", "의미"], triggers, col_widths=[4.0, 4.0, 9.0])

    add_heading(doc, "모드 B — 정기 검토 (매 시간 cron, hourly_check)", level=2)
    add_para(doc,
        "force_check=True 옵션. 트리거 평가/dedupe 모두 건너뛰고 monitor set 전종목에 대해 LLM 호출. "
        "트리거 못 잡은 \"slow drift\" 케이스 대응 — 가격이 천천히 무너지고 있는데 ±2σ는 안 찍힌 경우 등.",
        size=10)
    add_para(doc,
        "장점: 매 시간 LLM이 모든 보유 종목을 한 번씩 봐주니까 누락 없음. \n"
        "단점: 비용. 보유 5종목이면 시간당 5회 호출 × 7시간 = 일 35회 LLM 호출.",
        size=10)
    add_para(doc, "실행 명령: python -m scripts.daily_pipeline --phase intraday-loop --force",
             size=9, color=RGBColor(0x55, 0x55, 0x55))

    add_heading(doc, "비교", level=2)
    add_table(
        doc,
        headers=["", "트리거 기반 (15분)", "정기 검토 (매 시간)"],
        rows=[
            ("호출 시점",     "09:50/10:05/.../15:55 (15분 간격)", "09:30/10:30/.../15:30 (1시간 간격, 7회/일)"),
            ("LLM 호출 조건", "트리거 1개+ AND dedupe 통과",        "monitor set 전체 (조건 없음)"),
            ("dedupe",        "30분 안 중복 skip",                  "skip 안 함"),
            ("주요 용도",     "급격한 변화 감지 (spike/news)",     "slow drift / 정기 점검"),
            ("비용 (5종목 기준)", "트리거 충족 시만 — 일 5~15회",   "고정 35회/일"),
        ],
        col_widths=[3.5, 6.5, 7.0],
    )

    # ─── 18. LLM 결정 ───
    doc.add_page_break()
    add_heading(doc, "18. LLM 결정 — 5개 Action", level=1)
    add_para(doc, "Gemini는 다음 5개 중 하나를 선택합니다 (intraday_check 프롬프트):", size=10)

    actions = [
        ("enter",
         "신규 진입",
         "현재 보유 X일 때만 유효. position cap 5종목 / sector cap 2 / daily loss halt 안 걸려야. "
         "entry/stop/target_1r/target_2r 필수."),
        ("add",
         "포지션 추가",
         "이미 보유 중일 때만. 현재가 위로 add-on entry 지정. PEAD(어닝 직후 베타) 케이스에서 주로."),
        ("trim",
         "부분 매도",
         "이미 보유 중일 때. qty = 매도할 수량(보유의 일부). \"price_spike\"에서 이익 일부 확정용. "
         "Approve 시 50% 시장가 SELL로 자동 실행."),
        ("exit",
         "전량 청산",
         "이미 보유 중일 때. qty = 보유 전체. 부정적 뉴스 / stop 근처 fundamental 변화 시. "
         "Approve 시 close_position 자동 실행 (자식 미체결 주문 자동 cancel)."),
        ("hold",
         "유지 (no action)",
         "default. confidence < 0.6일 때 기본값. DB에도 저장 안 함 (액션 없음). "
         "단순 기록만 — Telegram 알림 X."),
    ]
    add_table(doc, ["Action", "의미", "조건/실행"], actions, col_widths=[2.5, 3.5, 11.0])

    add_heading(doc, "규칙 (시스템이 강제하는 조건)", level=2)
    add_bullet(doc, "regime.long_blocked = true → action은 hold 또는 exit만 허용")
    add_bullet(doc, "position이 null → add/trim 불가 (enter 또는 hold만)")
    add_bullet(doc, "position 있음 → enter 불가 (add/trim/exit/hold만)")
    add_bullet(doc, "confidence < 0.6 → Telegram 알림 안 보냄 (DB에는 저장)")
    add_bullet(doc, "5분 후 자동 expired (ADVISOR_APPROVAL_TTL_SEC) — 응답 못 받으면 만료")

    # ─── 19. 승인 실행 ───
    add_heading(doc, "19. 사용자 승인 → 자동 실행", level=1)
    add_para(doc,
        "Telegram inline 키보드에서 Approve 버튼을 누르면 webhook이 백엔드로 들어와서 즉시 broker API를 호출. "
        "rec_type별로 실행 경로가 다릅니다:",
        size=10)

    approve_paths = [
        ("intraday_entry / intraday_add",
         "trade_plan upsert (user_fixed 모드) + 즉시 dispatch_plan_immediately()로 bracket order 발송. "
         "amount_usd = equity / 5 (또는 사용자가 직접 입력)."),
        ("intraday_trim",
         "보유 수량의 50%를 시장가 SELL. broker.submit_order(MarketOrderRequest, qty = pos.qty // 2)."),
        ("intraday_exit",
         "broker.close_position(symbol) — 보유 전량 시장가 매도 + 미체결 자식(stop/target) 자동 cancel."),
        ("morning_brief (장 시작 전)",
         "위와 동일하게 trade_plan upsert + 즉시 발송. Phase 4 preopen에서 generate된 morning 추천 처리용."),
    ]
    add_table(doc, ["rec_type", "Approve 시 실행"], approve_paths, col_widths=[5.0, 12.0])

    add_para(doc, "Reject는 사유와 함께 advisor_recommendations.status='rejected'로 저장. "
                  "이 reject 데이터가 v10의 feedback_bonus와 auto_blacklist에 영향을 줌 — "
                  "Advisor가 reject한 종목은 다음 v10 픽에서 점수 차감.",
             size=10)

    # ─── 20. FAQ ───
    doc.add_page_break()
    add_heading(doc, "20. Advisor 자주 묻는 질문", level=1)
    qa_advisor = [
        ("Q. Advisor가 종목을 새로 발굴하나?",
         "A. 아님. monitor set = 오늘 발송된 plan + 현재 보유 포지션. 즉 v10/orb_auto가 이미 매수한 종목만 자문. "
         "예외: morning_brief (장 시작 전 자문)은 v10 picks Top 5에서 enter 추천 가능 — 하지만 그것도 v10 후보 풀 안에서만."),
        ("Q. \"매 시간 추천\"인데 진짜 매번 알림이 오나?",
         "A. 아님. LLM이 hold로 결정하면 알림 X. 또 confidence < 0.6이면 알림 X. "
         "실제로는 \"의미 있는 액션이 필요한 종목\"만 Telegram으로 옴 — 일 평균 1~3건 정도."),
        ("Q. 트리거 cron(15분)과 hourly cron(1시간) 차이가 뭔가?",
         "A. 15분 cron은 가격/뉴스/거래량 비정상 신호가 있을 때만 LLM 호출 — 비용 절약. "
         "1시간 cron은 무조건 모든 보유 종목 LLM 호출 — \"slow drift\" 케이스 잡기 위해. "
         "둘 다 동시에 운영하면 보완적 — 급변(15분) + 정기 점검(1시간) 둘 다 커버."),
        ("Q. LLM이 잘못된 판단해서 손실 나면?",
         "A. LLM 추천이 자동 실행되지 않음. 무조건 사용자 Approve 필요. "
         "사용자가 Telegram에서 \"Reject\"를 누르거나 5분간 응답 안 하면 자동 expired. "
         "또 reject 데이터가 누적되면 v10 auto_blacklist에 등록되어 시스템 self-correction."),
        ("Q. 어떤 LLM을 쓰나? 왜 Claude가 아니라 Gemini?",
         "A. Gemini 2.5 Flash. 비용 효율 때문 (Claude Opus의 1/20 가격). "
         "장중 매시간 호출에 Opus는 비용이 너무 큼. 정확도는 Flash로 충분 — \"이미 v10이 1차 필터한 후보\"라서 LLM은 \"이 종목 지금 액션 필요한가\"만 판단."),
        ("Q. confidence 임계값(0.6)을 더 높이면?",
         "A. 0.6은 ADVISOR_INTRADAY_MIN_CONFIDENCE 환경변수. 0.7로 올리면 알림이 줄고 정확도 향상, 단 놓치는 케이스 늘어남. "
         "0.5로 낮추면 알림 폭증 + 노이즈. 현재 0.6은 보수적 균형점."),
        ("Q. dedupe 30분이 짧지 않나? 같은 종목에 여러 번 추천이 오면?",
         "A. 30분 = DEDUPE_WINDOW_MIN. 너무 길게 하면 진짜 중요한 후속 신호를 놓침 (예: 첫 trim 후 30분 뒤 exit가 필요한 상황). "
         "정기 검토(force_check) 모드는 dedupe 무시 — 매 시간 항상 한 번씩 보장."),
    ]
    for q, a in qa_advisor:
        add_para(doc, q, size=10, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
        add_para(doc, a, size=10)

    # ─── 21. 통합 참고 ───
    add_heading(doc, "21. 참고 — 핵심 코드 위치", level=1)
    add_para(doc, "v10 (스윙)", size=11, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
    refs_v10 = [
        ("v10 알고리즘",      "scanner/integrated/run.py:208  (run_integrated_v10)"),
        ("v3 historical",     "scanner/comparison/v3_historical.py:93"),
        ("v3 게이트/점수",    "scanner/stage2_daily_picks.py:382 (evaluate_gates), :492 (evaluate_scores)"),
        ("scanner 후보",      "scanner/comparison/adapters.py (fetch_scanner_picks)"),
        ("entry/stop 산출",   "scanner/stage2_daily_picks.py:805 (compute_pick_metadata)"),
        ("trading 응답 빌더", "api/routes/trading.py:354 (_build_picks)"),
        ("점수 분해 (UI용)",  "api/routes/trading.py:271 (compute_score_breakdown)"),
    ]
    add_table(doc, headers=["역할", "파일:라인"], rows=refs_v10, col_widths=[5.0, 12.0])

    add_para(doc, "orb_auto (단타)", size=11, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
    refs_orb = [
        ("5-Model Stack (Phase 4)", "scanner/integrated/run.py:676 (run_integrated_intraday)"),
        ("Preopen cron",            "scripts/daily_pipeline.py:569 (run_preopen_phase)"),
        ("ORB 평가 (Phase 5)",      "scripts/intraday_confirm.py:127 (run_confirm)"),
        ("ORB 4-pass 신호",         "signals/opening_range.py:161 (evaluate_orb)"),
        ("entry/stop 산출 (ORB)",   "signals/opening_range.py:215 (compute_entry_stop_target)"),
        ("ORB 백테스트",            "backtests/run_intraday_orb.py"),
        ("백테스트 결과 파일",      "backtests/results/intraday_orb_*.json"),
    ]
    add_table(doc, headers=["역할", "파일:라인"], rows=refs_orb, col_widths=[5.0, 12.0])

    add_para(doc, "Advisor (AI 자문)", size=11, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
    refs_advisor = [
        ("Intraday 모니터 + 트리거",  "services/advisor/intraday_monitor.py:36 (run_intraday_loop_iteration)"),
        ("트리거 평가",                "services/advisor/intraday_monitor.py:139 (_evaluate_triggers)"),
        ("LLM 호출 + DB 저장",         "services/advisor/service.py:187 (run_intraday_check)"),
        ("Morning brief (장 시작 전)", "services/advisor/service.py:59 (run_morning_brief)"),
        ("승인 → broker 실행",         "services/advisor/service.py:413 (approve_recommendation)"),
        ("컨텍스트 빌더 (LLM 입력)",   "services/advisor/context_builder.py:43 (build_morning_context), :157 (build_intraday_context)"),
        ("LLM 클라이언트 (Gemini)",    "services/advisor/llm/google.py"),
        ("Intraday 프롬프트",          "services/advisor/prompts/v1/intraday_check.md"),
        ("Telegram 알림",              "services/advisor/notifications/telegram.py"),
        ("Cron 등록 (intraday-loop)",  "scripts/daily_pipeline.py:1274 (run_intraday_loop)"),
        ("API 라우트 (approve/reject)", "api/routes/advisor.py"),
    ]
    add_table(doc, headers=["역할", "파일:라인"], rows=refs_advisor, col_widths=[5.0, 12.0])

    add_para(doc, "메모리 참조 (관련 의사결정 기록)", size=11, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
    add_bullet(doc, "[[project_integrated_v10]] — v10 백테스트 10d alpha +14.88%, win 93%, Sharpe 6.58")
    add_bullet(doc, "[[project_hybrid_dispatch]] — 2026-05-13 도입, user_fixed(09:30) + orb_auto(09:45) 분리")
    add_bullet(doc, "[[project_2tier_partial_exit]] — 50:50 분할 (1차 +1R, 2차 +2R), 2026-05-08 도입")
    add_bullet(doc, "[[project_autotrade_v1_production]] — 2026-05-09 paper production-ready, 9중 안전장치")
    add_bullet(doc, "[[project_bracket_protection_fix]] — TIF day→gtc, 보호 stop 자동 재발송")
    add_bullet(doc, "[[project_advisor_v1]] — Gemini 2.5 Flash, 매시간 cron + Telegram webhook, Approve 즉시 발송")
    add_bullet(doc, "[[project_activity_and_alerts]] — /activity 시간순 통합 timeline (7소스), heartbeat Telegram alert")

    add_para(doc, "— 끝 —", size=9, color=RGBColor(0x80, 0x80, 0x80))

    return doc


if __name__ == "__main__":
    out_dir = Path(__file__).parent.parent / "docs"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / "v10_종목선정_설명서.docx"
    doc = build_doc()
    doc.save(str(out_path))
    import sys
    msg = f"saved: {out_path}\nsize: {out_path.stat().st_size:,} bytes"
    try:
        print(msg)
    except UnicodeEncodeError:
        sys.stdout.buffer.write(msg.encode("utf-8") + b"\n")
